import os
import sys
import webbrowser
import ctypes
import traceback

# Support Unicode Vietnamese characters in Windows Console
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

# Handle dependencies imports safely and show message if missing
try:
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build
    from docx import Document
    import pandas as pd
except ImportError as e:
    ctypes.windll.user32.MessageBoxW(
        0,
        f"Lỗi: Thiếu thư viện Python cần thiết.\nChi tiết: {e}\n\nVui lòng cài đặt bằng cách chạy file setup hoặc lệnh 'pip install -r requirements.txt'.",
        "Lỗi Thư Viện",
        0x10
    )
    sys.exit(1)

# Google API Scopes
SCOPES = [
    'https://www.googleapis.com/auth/spreadsheets',
    'https://www.googleapis.com/auth/drive.file'
]

def show_message(title, text, icon_type=0):
    # 0 = Info, 0x10 = Error, 0x30 = Warning, 0x40 = Info
    return ctypes.windll.user32.MessageBoxW(0, text, title, icon_type)

def get_credentials():
    creds = None
    credentials_path, app_dir = resolve_credentials_path()
    token_path = os.path.join(app_dir, 'token.json')

    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception:
                creds = None

        if not creds:
            if not os.path.exists(credentials_path):
                msg = (
                    "Không tìm thấy file cấu hình xác thực Google (credentials.json) trong thư mục ứng dụng.\n\n"
                    "Vui lòng:\n"
                    "1. Tạo dự án trên Google Cloud Console.\n"
                    "2. Tạo OAuth 2.0 Client ID (chọn Application type là Desktop App).\n"
                    "3. Tải file JSON cấu hình về, đổi tên thành 'credentials.json' và copy vào thư mục:\n"
                    f"'{app_dir}'\n\n"
                    "Sau đó bấm chuột phải thử lại."
                )
                show_message("Thiếu cấu hình Google Credentials", msg, 0x10)
                sys.exit(1)

            print("Đang mở trình duyệt để xác thực tài khoản Google...")
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)

            import json
            try:
                with open(credentials_path, 'r', encoding='utf-8') as f:
                    cred_data = json.load(f)
                if 'web' in cred_data:
                    redirect_uris = cred_data['web'].get('redirect_uris', [])
                    local_redirect = None
                    for uri in redirect_uris:
                        if 'localhost:8900' in uri:
                            local_redirect = uri
                            break
                    if local_redirect:
                        flow.redirect_uri = local_redirect
                        print("Phát hiện Web Client credentials. Sử dụng callback port 8900...")
                        creds = flow.run_local_server(port=8900)
                    else:
                        creds = flow.run_local_server(port=0)
                else:
                    creds = flow.run_local_server(port=0)
            except Exception as e:
                print(f"Lỗi khi đọc file credentials: {e}")
                creds = flow.run_local_server(port=0)

        with open(token_path, 'w') as token:
            token.write(creds.to_json())

    return creds

    raise last_error

def parse_hex_color(hex_str):
    """Converts hex color (e.g. FFFFFFFF, FF0000, 000000) to RGB floats 0.0-1.0"""
    if not hex_str or not isinstance(hex_str, str):
        return None

    # Strip alpha if 8 characters
    if len(hex_str) == 8:
        # e.g., 'FFFF0000' -> alpha is 'FF', RGB is 'FF0000'
        # If alpha is 00 (fully transparent), treat as transparent / no fill
        if hex_str.startswith('00'):
            return None
        hex_str = hex_str[2:]

    if len(hex_str) != 6:
        return None

    try:
        r = int(hex_str[0:2], 16) / 255.0
        g = int(hex_str[2:4], 16) / 255.0
        b = int(hex_str[4:6], 16) / 255.0
        return {"red": r, "green": g, "blue": b}
    except ValueError:
        return None

def get_cell_format(cell):
    cell_format = {}

    # 1. Background color
    if cell.fill and cell.fill.fill_type == 'solid':
        if cell.fill.start_color and cell.fill.start_color.type == 'rgb':
            bg = parse_hex_color(cell.fill.start_color.rgb)
            if bg:
                cell_format["backgroundColor"] = bg

    # 2. Font properties
    font_format = {}
    if cell.font:
        if cell.font.name:
            font_format["fontFamily"] = cell.font.name
        if cell.font.size:
            font_format["fontSize"] = int(cell.font.size)
        if cell.font.bold:
            font_format["bold"] = True
        if cell.font.italic:
            font_format["italic"] = True

        # Text color
        if cell.font.color and cell.font.color.type == 'rgb':
            fg = parse_hex_color(cell.font.color.rgb)
            if fg:
                font_format["foregroundColor"] = fg

    if font_format:
        cell_format["textFormat"] = font_format

    # 3. Alignment
    if cell.alignment:
        h_align = cell.alignment.horizontal
        if h_align in ['left', 'center', 'right', 'justify']:
            cell_format["horizontalAlignment"] = h_align.upper()
        v_align = cell.alignment.vertical
        if v_align == 'center':
            cell_format["verticalAlignment"] = 'MIDDLE'
        elif v_align in ['top', 'bottom']:
            cell_format["verticalAlignment"] = v_align.upper()

        # 4. Wrap text
        if cell.alignment.wrap_text:
            cell_format["wrapStrategy"] = "WRAP"

    return cell_format

def parse_docx(file_path):
    doc = Document(file_path)
    sheets_data = {}

    # 1. Parse tables
    has_tables = False
    for i, table in enumerate(doc.tables):
        has_tables = True
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
        if table_rows:
            sheets_data[f"Table {i+1}"] = {
                'values': table_rows,
                'styles': None,
                'col_widths': None,
                'row_heights': None
            }

    # 2. Parse paragraphs (text content)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        # Format paragraphs as a single-column table
        paragraphs_data = [[p] for p in paragraphs]
        if has_tables:
            sheets_data["Text Content"] = {
                'values': paragraphs_data,
                'styles': None,
                'col_widths': None,
                'row_heights': None
            }
        else:
            sheets_data["Sheet1"] = {
                'values': paragraphs_data,
                'styles': None,
                'col_widths': None,
                'row_heights': None
            }

    if not sheets_data:
        sheets_data["Sheet1"] = {
            'values': [["File Word này không có nội dung văn bản hoặc bảng biểu."]],
            'styles': None,
            'col_widths': None,
            'row_heights': None
        }

    return sheets_data

def parse_spreadsheet(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    sheets_data = {}

    if ext == '.csv':
        df = read_csv_with_fallback(file_path)
        df = df.fillna('')
        values = [df.columns.tolist()] + df.values.tolist()
        sheets_data['CSV Data'] = {
            'values': values,
            'styles': None,
            'col_widths': None,
            'row_heights': None
        }
    elif ext == '.xlsx':
        import openpyxl
        from openpyxl.utils import get_column_letter

        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            values = []
            styles = []

            # Read row heights
            row_heights = {}
            for r_idx in range(1, ws.max_row + 1):
                dim = ws.row_dimensions.get(r_idx)
                if dim and dim.height:
                    px = int(dim.height * 1.3)
                    if 10 <= px <= 500:
                        row_heights[r_idx - 1] = px

            # Read col widths
            col_widths = {}
            for col_idx in range(1, ws.max_column + 1):
                col_letter = get_column_letter(col_idx)
                dim = ws.column_dimensions.get(col_letter)
                if dim and dim.width:
                    px = int(dim.width * 8)
                    if 20 <= px <= 1000:
                        col_widths[col_idx - 1] = px

            # Read values and styles
            for r_idx in range(1, ws.max_row + 1):
                row_vals = []
                row_styles = []
                for c_idx in range(1, ws.max_column + 1):
                    cell = ws.cell(row=r_idx, column=c_idx)
                    val = cell.value
                    if val is None:
                        val = ''
                    row_vals.append(val)
                    row_styles.append(get_cell_format(cell))
                values.append(row_vals)
                styles.append(row_styles)

            safe_name = sheet_name[:30]
            sheets_data[safe_name] = {
                'values': values,
                'styles': styles,
                'col_widths': col_widths,
                'row_heights': row_heights
            }
    else:
        # Excel legacy (.xls)
        xl = pd.ExcelFile(file_path)
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet_name)
            df = df.fillna('')
            values = [df.columns.tolist()] + df.values.tolist()
            safe_name = sheet_name[:30]
            sheets_data[safe_name] = {
                'values': values,
                'styles': None,
                'col_widths': None,
                'row_heights': None
            }

    return sheets_data

def style_spreadsheet(service, spreadsheet_id, sheet_mapping, sheets_data):
    """
    Applies styling to the Google Spreadsheet.
    - If sheet has parsed styles (Excel), applies them directly.
    - Otherwise, applies default professional styling.
    """
    try:
        print("Đang định dạng trang tính (styling)...")
        requests = []
        for sheet_name, info in sheets_data.items():
            sheet_id = sheet_mapping.get(sheet_name)
            if sheet_id is None:
                continue

            values = info.get('values', [])
            if not values:
                continue

            row_count = len(values)
            col_count = max(len(row) for row in values) if row_count > 0 else 0
            if row_count == 0 or col_count == 0:
                continue

            # 1. Enable grid lines (hideGridlines = False)
            requests.append({
                "updateSheetProperties": {
                    "properties": {
                        "sheetId": sheet_id,
                        "gridProperties": {
                            "hideGridlines": False
                        }
                    },
                    "fields": "gridProperties.hideGridlines"
                }
            })

            excel_styles = info.get('styles')
            col_widths = info.get('col_widths')
            row_heights = info.get('row_heights')

            # Case A: Apply styles parsed from Excel
            if excel_styles:
                rows_data = []
                for r_idx in range(row_count):
                    row_cells = []
                    for c_idx in range(col_count):
                        cell_fmt = {}
                        if r_idx < len(excel_styles) and c_idx < len(excel_styles[r_idx]):
                            cell_fmt = excel_styles[r_idx][c_idx]
                        row_cells.append({
                            "userEnteredFormat": cell_fmt
                        })
                    rows_data.append({
                        "values": row_cells
                    })

                requests.append({
                    "updateCells": {
                        "range": {
                            "sheetId": sheet_id,
                            "startRowIndex": 0,
                            "endRowIndex": row_count,
                            "startColumnIndex": 0,
                            "endColumnIndex": col_count
                        },
                        "rows": rows_data,
                        "fields": "userEnteredFormat"
                    }
                })

                # Apply column widths
                if col_widths:
                    for c_idx, width in col_widths.items():
                        if c_idx < col_count:
                            requests.append({
                                "updateDimensionProperties": {
                                    "range": {
                                        "sheetId": sheet_id,
                                        "dimension": "COLUMNS",
                                        "startIndex": c_idx,
                                        "endIndex": c_idx + 1
                                    },
                                    "properties": {
                                        "pixelSize": width
                                    },
                                    "fields": "pixelSize"
                                }
                            })
                else:
                    requests.append({
                        "autoResizeDimensions": {
                            "dimensions": {
                                "sheetId": sheet_id,
                                "dimension": "COLUMNS",
                                "startIndex": 0,
                                "endIndex": col_count
                            }
                        }
                    })

                # Apply row heights
                if row_heights:
                    for r_idx, height in row_heights.items():
                        if r_idx < row_count:
                            requests.append({
                                "updateDimensionProperties": {
                                    "range": {
                                        "sheetId": sheet_id,
                                        "dimension": "ROWS",
                                        "startIndex": r_idx,
                                        "endIndex": r_idx + 1
                                    },
                                    "properties": {
                                        "pixelSize": height
                                    },
                                    "fields": "pixelSize"
                                }
                            })

            # Case B: Standard fallback style (Word, CSV, legacy XLS)
            else:
                # Set Row Heights: Header row (index 0) to 40px, Data rows (index 1+) to 28px
                requests.append({
                    "updateDimensionProperties": {
                        "range": {
                            "sheetId": sheet_id,
                            "dimension": "ROWS",
                            "startIndex": 0,
                            "endIndex": 1
                        },
                        "properties": {
                            "pixelSize": 40
                        },
                        "fields": "pixelSize"
                    }
                })

                if row_count > 1:
                    requests.append({
                        "updateDimensionProperties": {
                            "range": {
                                "sheetId": sheet_id,
                                "dimension": "ROWS",
                                "startIndex": 1,
                                "endIndex": row_count
                            },
                            "properties": {
                                "pixelSize": 28
                            },
                            "fields": "pixelSize"
                        }
                    })

                # Format header cells (Row 0): Google Blue Bg, White Text, Bold, Centered
                requests.append({
                    "repeatCell": {
                        "range": {
                            "sheetId": sheet_id,
                            "startRowIndex": 0,
                            "endRowIndex": 1,
                            "startColumnIndex": 0,
                            "endColumnIndex": col_count
                        },
                        "cell": {
                            "userEnteredFormat": {
                                "backgroundColor": {
                                    "red": 26/255.0,
                                    "green": 115/255.0,
                                    "blue": 232/255.0
                                },
                                "textFormat": {
                                    "foregroundColor": {
                                        "red": 1.0,
                                        "green": 1.0,
                                        "blue": 1.0
                                    },
                                    "fontFamily": "Segoe UI",
                                    "fontSize": 11,
                                    "bold": True
                                },
                                "horizontalAlignment": "CENTER",
                                "verticalAlignment": "MIDDLE"
                            }
                        },
                        "fields": "userEnteredFormat(backgroundColor,textFormat,horizontalAlignment,verticalAlignment)"
                    }
                })

                # Format data cells general defaults: Segoe UI, 10pt, Vertical Middle
                if row_count > 1:
                    requests.append({
                        "repeatCell": {
                            "range": {
                                "sheetId": sheet_id,
                                "startRowIndex": 1,
                                "endRowIndex": row_count,
                                "startColumnIndex": 0,
                                "endColumnIndex": col_count
                            },
                            "cell": {
                                "userEnteredFormat": {
                                    "textFormat": {
                                        "fontFamily": "Segoe UI",
                                        "fontSize": 10
                                    },
                                    "verticalAlignment": "MIDDLE",
                                    "wrapStrategy": "WRAP"
                                }
                            },
                            "fields": "userEnteredFormat(textFormat,verticalAlignment,wrapStrategy)"
                        }
                    })

                    # Apply Alternating (Zebra) Row background color to data rows
                    for r_idx in range(1, row_count):
                        bg_color = {
                            "red": 248/255.0,
                            "green": 249/255.0,
                            "blue": 250/255.0
                        } if r_idx % 2 == 1 else {
                            "red": 1.0,
                            "green": 1.0,
                            "blue": 1.0
                        }

                        requests.append({
                            "repeatCell": {
                                "range": {
                                    "sheetId": sheet_id,
                                    "startRowIndex": r_idx,
                                    "endRowIndex": r_idx + 1,
                                    "startColumnIndex": 0,
                                    "endColumnIndex": col_count
                                },
                                "cell": {
                                    "userEnteredFormat": {
                                        "backgroundColor": bg_color
                                    }
                                },
                                "fields": "userEnteredFormat(backgroundColor)"
                            }
                        })

                # Apply subtle gray borders around all populated cells
                border_color = {"red": 218/255.0, "green": 220/255.0, "blue": 224/255.0} # #dadce0
                requests.append({
                    "updateBorders": {
                        "range": {
                            "sheetId": sheet_id,
                            "startRowIndex": 0,
                            "endRowIndex": row_count,
                            "startColumnIndex": 0,
                            "endColumnIndex": col_count
                        },
                        "top": {"style": "SOLID", "color": border_color},
                        "bottom": {"style": "SOLID", "color": border_color},
                        "left": {"style": "SOLID", "color": border_color},
                        "right": {"style": "SOLID", "color": border_color},
                        "innerHorizontal": {"style": "SOLID", "color": border_color},
                        "innerVertical": {"style": "SOLID", "color": border_color}
                    }
                })

                # Auto-resize all columns
                requests.append({
                    "autoResizeDimensions": {
                        "dimensions": {
                            "sheetId": sheet_id,
                            "dimension": "COLUMNS",
                            "startIndex": 0,
                            "endIndex": col_count
                        }
                    }
                })

        if requests:
            body = {"requests": requests}
            service.spreadsheets().batchUpdate(spreadsheetId=spreadsheet_id, body=body).execute()
            print("Đã định dạng trang tính thành công!")
    except Exception as e:
        print(f"Lỗi khi điịnh dạng trang tính: {e}")

def main():
    if len(sys.argv) < 2:
        show_message("Lỗi Tham Số", "Ứng dụng được gọi mà không có đường dẫn file. Vui lòng sử dụng qua menu chuột phải của Windows.", 0x10)
        sys.exit(1)

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        show_message("Lỗi File", f"Không tìm thấy file: {file_path}", 0x10)
        sys.exit(1)

    filename = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    print(f"Đang chuẩn bị import file: {filename}...")

    # 1. Parse file content
    try:
        if ext == '.docx':
            sheets_data = parse_docx(file_path)
        elif ext in ['.xlsx', '.xls', '.csv']:
            sheets_data = parse_spreadsheet(file_path)
        else:
            show_message("Định Dạng Không Hỗ Trợ", f"Không hỗ trợ định dạng file: {ext}", 0x10)
            sys.exit(1)
    except Exception as e:
        show_message("Lỗi Đọc File", f"Có lỗi xảy ra khi đọc file:\n{e}", 0x10)
        traceback.print_exc()
        sys.exit(1)

    # 2. Authenticate Google API
    try:
        creds = get_credentials()
        service = build('sheets', 'v4', credentials=creds)
    except Exception as e:
        show_message("Lỗi Xác Thực", f"Không thể kết nối đến Google API:\n{e}", 0x10)
        traceback.print_exc()
        sys.exit(1)

    # 3. Create Google Spreadsheet
    try:
        print("Đang tạo Google Spreadsheet mới...")
        sheet_titles = list(sheets_data.keys())

        # Prepare body with sheets definitions
        sheets_def = [{'properties': {'title': title}} for title in sheet_titles]

        body = {
            'properties': {
                'title': f"[Imported] {os.path.splitext(filename)[0]}"
            },
            'sheets': sheets_def
        }

        spreadsheet = service.spreadsheets().create(
            body=body,
            fields='spreadsheetId,spreadsheetUrl,sheets(properties(title,sheetId))'
        ).execute()
        spreadsheet_id = spreadsheet.get('spreadsheetId')
        spreadsheet_url = spreadsheet.get('spreadsheetUrl')

        # Map sheet title to sheetId
        sheet_mapping = {}
        for s in spreadsheet.get('sheets', []):
            prop = s.get('properties', {})
            sheet_mapping[prop.get('title')] = prop.get('sheetId')

        # 4. Populate Spreadsheet content
        print("Đang tải dữ liệu lên Google Sheets...")
        data = []
        for sheet_name, info in sheets_data.items():
            data.append({
                'range': f"'{sheet_name}'!A1",
                'values': info['values']
            })

        body_update = {
            'valueInputOption': 'USER_ENTERED',
            'data': data
        }

        service.spreadsheets().values().batchUpdate(spreadsheetId=spreadsheet_id, body=body_update).execute()
        print("Đã tải dữ liệu thô lên thành công.")

        # 5. Apply styling
        style_spreadsheet(service, spreadsheet_id, sheet_mapping, sheets_data)

        print(f"Import thành công! Link Google Sheet: {spreadsheet_url}")

        # Auto-open in browser
        webbrowser.open(spreadsheet_url)

        # Show success toast/notification
        show_message("Import Thành Công", f"Đã import thành công file '{filename}' vào Google Sheets và đang mở trên trình duyệt của bạn.", 0x40)

    except Exception as e:
        show_message("Lỗi Upload Google Sheets", f"Gặp lỗi khi tạo/ghi vào Google Sheets:\n{e}", 0x10)
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        show_message("Lỗi Hệ Thống", f"Đã xảy ra lỗi không mong muốn:\n{e}", 0x10)
        traceback.print_exc()
        input("\nNhấn Enter để đóng cửa sổ...")
